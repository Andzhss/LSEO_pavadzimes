from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import io
import os

# --- Krāsas (sky blue theme) ---
COLOR_DARK  = "1A6FA8"   # tumšzils — galvenes
COLOR_MID   = "4DA6E8"   # vidējs zils — apmales
COLOR_LIGHT = "D0EAF8"   # gaiši zils — rindu fons
COLOR_BG    = "EAF4FB"   # ļoti gaiši zils — bloku foni
COLOR_WHITE = "FFFFFF"

# --- Tulkojumi DOCX dokumentam ---
DOCX_TRANSLATIONS = {
    "lv": {
        "client":          "KLIENTS",
        "sender":          "PIEGĀDĀTĀJS",
        "receiver":        "Saņēmējs",
        "customer":        "Pasūtītājs",
        "sender_label":    "Nosūtītājs",
        "address":         "Adrese",
        "reg_no":          "Reģ. Nr.",
        "vat_no":          "PVN Nr.",
        "phone":           "Tālrunis",
        "swift":           "SWIFT/BIC",
        "account":         "Bankas konta numurs",
        "due_date":        "Apmaksāt līdz",
        "date":            "Datums",
        "col_name":        "NOSAUKUMS",
        "col_unit":        "Mērvienība",
        "col_qty":         "DAUDZUMS",
        "col_price":       "CENA (EUR)",
        "col_total":       "KOPĀ (EUR)",
        "total_label":     "KOPĀ",
        "total_no_vat":    "KOPĀ (bez PVN un atlaides)",
        "discount":        "Atlaides apjoms",
        "total_discount":  "Kopā ar atlaidi (bez PVN)",
        "vat":             "PVN",
        "grand_total":     "Kopumā",
        "advance_payable": "APMAKSĀJAMAIS AVANSS",
        "words_prefix":    "Vārdiem: ",
        "extra_info":      "Papildus informācija:",
        "prep_invoice":    "Pavadzīmi sagatavoja: ",
        "prep_receipt":    "Rēķinu sagatavoja: ",
        "prep_advance":    "Avansa rēķinu sagatavoja: ",
        "recv_invoice":    "Pavadzīmi saņēma:",
        "recv_receipt":    "Rēķinu saņēma:",
        "recv_advance":    "Avansa rēķinu saņēma:",
    },
    "en": {
        "client":          "CLIENT",
        "sender":          "SUPPLIER",
        "receiver":        "Receiver",
        "customer":        "Customer",
        "sender_label":    "Sender",
        "address":         "Address",
        "reg_no":          "Reg. No.",
        "vat_no":          "VAT No.",
        "phone":           "Phone",
        "swift":           "SWIFT/BIC",
        "account":         "Bank account number",
        "due_date":        "Due date",
        "date":            "Date",
        "col_name":        "DESCRIPTION",
        "col_unit":        "Unit",
        "col_qty":         "QUANTITY",
        "col_price":       "PRICE (EUR)",
        "col_total":       "TOTAL (EUR)",
        "total_label":     "SUBTOTAL",
        "total_no_vat":    "SUBTOTAL (excl. VAT and discount)",
        "discount":        "Discount amount",
        "total_discount":  "Total with discount (excl. VAT)",
        "vat":             "VAT",
        "grand_total":     "Total",
        "advance_payable": "ADVANCE PAYABLE",
        "words_prefix":    "In words: ",
        "extra_info":      "Additional information:",
        "prep_invoice":    "Invoice prepared by: ",
        "prep_receipt":    "Receipt prepared by: ",
        "prep_advance":    "Advance invoice prepared by: ",
        "recv_invoice":    "Invoice received by:",
        "recv_receipt":    "Receipt received by:",
        "recv_advance":    "Advance invoice received by:",
    }
}

def fmt_curr(val):
    return f"{val:,.2f}".replace(",", "X").replace(".", ",").replace("X", " ")

def _rgb(hex_str):
    r = int(hex_str[0:2], 16)
    g = int(hex_str[2:4], 16)
    b = int(hex_str[4:6], 16)
    return RGBColor(r, g, b)

def _set_cell_bg(cell, hex_color):
    """Iestata šūnas fona krāsu."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    tcPr.append(shd)

def _set_cell_borders(cell, color="4DA6E8", size="4"):
    """Iestata šūnas apmales."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top    w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>
            <w:left   w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>
            <w:right  w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)

def _set_row_height(row, height_cm):
    """Iestata rindas augstumu."""
    trPr = row._tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_cm * 567)}" w:hRule="atLeast"/>'
    )
    trPr.append(trHeight)

def _add_color_bar(doc, hex_color=COLOR_DARK, height_cm=0.22):
    """Pievieno krāsainu horizontālu joslu kā 1x1 tabulu."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(17)
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, hex_color)
    _set_row_height(tbl.rows[0], height_cm)
    # Notīra paragrāfu šūnā lai nav papildu vietas
    for p in cell.paragraphs:
        p.clear()

def _add_thin_line(doc, hex_color=COLOR_LIGHT):
    """Pievieno plānu horizontālu līniju."""
    _add_color_bar(doc, hex_color, height_cm=0.05)

def generate_docx(data):
    lang = data.get('lang', 'lv')
    tr   = DOCX_TRANSLATIONS.get(lang, DOCX_TRANSLATIONS['lv'])

    doc = Document()

    # --- Apmales ---
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # --- Noklusējuma fonts ---
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    # Dati
    doc_type     = data.get('doc_type', 'Pavadzīme')
    doc_id       = data.get('doc_id', 'LSEO 0001')
    date_str     = data.get('date', '')
    due_str      = data.get('due_date', '')
    is_e_invoice = "e-rēķins" in doc_type.lower() or "e-invoice" in doc_type.lower()
    is_advance   = "avansa" in doc_type.lower() or "advance" in doc_type.lower()
    display_type = "Invoice" if is_e_invoice and lang == "en" else ("Rēķins" if is_e_invoice else doc_type)

    current_dir = os.path.dirname(os.path.abspath(__file__))
    logo_path   = os.path.join(current_dir, "LatSEO logo black.png")

    # ==========================================
    # 1. VIRSRAKSTS
    # ==========================================
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.autofit = False
    hdr_table.columns[0].width = Cm(8.5)
    hdr_table.columns[1].width = Cm(8.5)

    # Logo
    cell_logo = hdr_table.cell(0, 0)
    try:
        cell_logo.paragraphs[0].add_run().add_picture(logo_path, width=Cm(4.0))
    except Exception:
        r = cell_logo.paragraphs[0].add_run("LatSEO")
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = _rgb(COLOR_DARK)

    # Dokumenta info
    cell_info = hdr_table.cell(0, 1)
    p = cell_info.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run_title = p.add_run(f"{display_type} Nr. {doc_id}\n")
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = _rgb(COLOR_DARK)

    p.add_run(f"\n{tr['date']}: ")
    run_date = p.add_run(f"{date_str}\n")
    run_date.bold = True

    p.add_run(f"{tr['due_date']}: ")
    run_due = p.add_run(due_str)
    run_due.bold = True

    doc.add_paragraph()
    _add_color_bar(doc, COLOR_DARK)
    _add_thin_line(doc, COLOR_LIGHT)

    # ==========================================
    # 2. KLIENTS UN PIEGĀDĀTĀJS / E-RĒĶINS
    # ==========================================
    doc.add_paragraph()

    if is_e_invoice:
        info_table = doc.add_table(rows=1, cols=2)
        info_table.autofit = False
        info_table.columns[0].width = Cm(8.3)
        info_table.columns[1].width = Cm(8.3)

        def _fill_e_cell(cell, title, name, reg, addr):
            _set_cell_bg(cell, COLOR_BG)
            _set_cell_borders(cell, COLOR_MID, "4")
            p = cell.paragraphs[0]
            r = p.add_run(f"{title}\n\n")
            r.bold = True
            r.font.color.rgb = _rgb(COLOR_DARK)
            r2 = p.add_run(f"{name}\n")
            r2.bold = True
            p.add_run(f"{tr['reg_no']}: {reg}\n").italic = True
            p.add_run(f"{tr['address']}: {addr}").italic = True

        _fill_e_cell(info_table.cell(0, 0),
            tr['receiver'], data.get('receiver_name',''),
            data.get('receiver_reg_no',''), data.get('receiver_address',''))
        _fill_e_cell(info_table.cell(0, 1),
            tr['customer'], data.get('customer_name',''),
            data.get('customer_reg_no',''), data.get('customer_address',''))

        doc.add_paragraph()
        _add_thin_line(doc, COLOR_LIGHT)
        doc.add_paragraph()

        p_sender = doc.add_paragraph()
        r = p_sender.add_run(tr['sender_label'])
        r.bold = True
        r.font.color.rgb = _rgb(COLOR_DARK)
    else:
        info_table = doc.add_table(rows=1, cols=2)
        info_table.autofit = False
        info_table.columns[0].width = Cm(8.3)
        info_table.columns[1].width = Cm(8.3)

        # Klients
        cell_client = info_table.cell(0, 0)
        _set_cell_bg(cell_client, COLOR_BG)
        _set_cell_borders(cell_client, COLOR_MID, "4")
        p_c = cell_client.paragraphs[0]
        r = p_c.add_run(f"{tr['client']}\n\n")
        r.bold = True
        r.font.color.rgb = _rgb(COLOR_DARK)
        r2 = p_c.add_run(f"{data.get('client_name','')}\n")
        r2.bold = True
        p_c.add_run(f"{tr['address']}: {data.get('client_address','')}\n").italic = True
        p_c.add_run(f"{tr['reg_no']}: {data.get('client_reg_no','')}\n").italic = True
        p_c.add_run(f"{tr['vat_no']}: {data.get('client_vat_no','')}").italic = True

        # Piegādātājs
        cell_sender = info_table.cell(0, 1)
        _set_cell_bg(cell_sender, COLOR_LIGHT)
        _set_cell_borders(cell_sender, COLOR_MID, "4")
        p_s = cell_sender.paragraphs[0]
        r = p_s.add_run(f"{tr['sender']}\n\n")
        r.bold = True
        r.font.color.rgb = _rgb(COLOR_DARK)
        r2 = p_s.add_run("SIA Baltic SEO\n")
        r2.bold = True
        p_s.add_run(f"{tr['address']}: Ķekavas nov., Ķekavas pag.,\nOdukalns, Kārklu iela 4, LV-2123\n").italic = True
        p_s.add_run(f"{tr['reg_no']}: 40203749304\n").italic = True
        p_s.add_run(f"{tr['phone']}: +371 24424434").italic = True

    # ==========================================
    # 3. BANKAS INFO
    # ==========================================
    doc.add_paragraph()

    bank_table = doc.add_table(rows=1, cols=4)
    bank_table.autofit = False
    bank_table.columns[0].width = Cm(3.0)
    bank_table.columns[1].width = Cm(3.5)
    bank_table.columns[2].width = Cm(4.0)
    bank_table.columns[3].width = Cm(6.5)

    bank_cells = [
        ("AS Swedbank", True),
        (f"{tr['swift']}: HABALV22", False),
        (f"{tr['account']}:", False),
        ("LV05HABA0551065262400", True),
    ]
    for i, (txt, bold) in enumerate(bank_cells):
        cell = bank_table.cell(0, i)
        _set_cell_bg(cell, COLOR_BG)
        _set_cell_borders(cell, COLOR_MID, "3")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(txt)
        run.bold = bold
        run.font.size = Pt(8.5)
        if bold:
            run.font.color.rgb = _rgb(COLOR_DARK)

    doc.add_paragraph()

    # ==========================================
    # 4. PREČU TABULA
    # ==========================================
    items = data.get('items', [])
    headers = [tr['col_name'], tr['col_unit'], tr['col_qty'], tr['col_price'], tr['col_total']]

    prod_table = doc.add_table(rows=1 + len(items), cols=5)
    prod_table.style = 'Table Grid'
    prod_table.columns[0].width = Cm(6.5)
    prod_table.columns[1].width = Cm(2.5)
    prod_table.columns[2].width = Cm(2.5)
    prod_table.columns[3].width = Cm(2.5)
    prod_table.columns[4].width = Cm(3.0)

    # Galvene
    hdr_cells = prod_table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        _set_cell_bg(cell, COLOR_DARK)
        _set_cell_borders(cell, COLOR_WHITE, "4")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = _rgb(COLOR_WHITE)

    # Datu rindas
    for row_i, item in enumerate(items):
        row = prod_table.rows[row_i + 1]
        bg  = COLOR_BG if row_i % 2 == 0 else COLOR_WHITE
        seq_num  = item.get('seq', '')
        name_str = f"{seq_num}. {item['name']}" if seq_num else item['name']

        values = [
            (name_str,            WD_ALIGN_PARAGRAPH.LEFT,   False),
            (item['unit'],        WD_ALIGN_PARAGRAPH.CENTER, False),
            (str(item['qty']),    WD_ALIGN_PARAGRAPH.CENTER, False),
            (str(item['price']),  WD_ALIGN_PARAGRAPH.RIGHT,  False),
            (str(item['total']),  WD_ALIGN_PARAGRAPH.RIGHT,  True),
        ]
        for col_i, (val, align, bold) in enumerate(values):
            cell = row.cells[col_i]
            _set_cell_bg(cell, bg)
            _set_cell_borders(cell, COLOR_MID, "3")
            p = cell.paragraphs[0]
            p.alignment = align
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if bold:
                run.bold = True

    doc.add_paragraph()

    # ==========================================
    # 5. KOPSUMMAS
    # ==========================================
    subtotal_val     = data.get('subtotal', '0,00')
    vat_val          = data.get('vat', '0,00')
    total_val        = data.get('total', '0,00')
    raw_discount_eur = data.get('raw_discount_eur', 0.0)
    apply_vat        = data.get('apply_vat', True)

    if raw_discount_eur > 0:
        sum_rows = [
            (tr['total_no_vat'],                                         subtotal_val,                           False),
            (f"{tr['discount']} ({data.get('discount_percent',0):g}%)", f"-{data.get('discount_eur','0,00')}",  False),
            (tr['total_discount'],                                        data.get('subtotal_after_discount','0,00'), False),
        ]
        if apply_vat:
            sum_rows.append((tr['vat'], vat_val, False))
        sum_rows.append((tr['grand_total'], total_val, True))
    else:
        sum_rows = [(tr['total_label'], subtotal_val, False)]
        if apply_vat:
            sum_rows.append((tr['vat'], vat_val, False))
        sum_rows.append((tr['grand_total'], total_val, True))

    sum_table = doc.add_table(rows=len(sum_rows), cols=3)
    sum_table.autofit = False
    sum_table.columns[0].width = Cm(8.5)
    sum_table.columns[1].width = Cm(5.0)
    sum_table.columns[2].width = Cm(3.5)

    for r_i, (label, value, is_bold) in enumerate(sum_rows):
        row = sum_table.rows[r_i]
        is_last = (r_i == len(sum_rows) - 1)
        bg = COLOR_LIGHT if is_last else COLOR_WHITE

        # tukša pirmā šūna
        _set_cell_bg(row.cells[0], bg)

        label_cell = row.cells[1]
        _set_cell_bg(label_cell, bg)
        if is_last:
            _set_cell_borders(label_cell, COLOR_DARK, "6")
        p = label_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{label}")
        run.bold = is_bold
        run.font.size = Pt(9.5 if is_bold else 9)
        if is_bold:
            run.font.color.rgb = _rgb(COLOR_DARK)

        val_cell = row.cells[2]
        _set_cell_bg(val_cell, bg)
        if is_last:
            _set_cell_borders(val_cell, COLOR_DARK, "6")
        p = val_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{value} €")
        run.bold = is_bold
        run.font.size = Pt(9.5 if is_bold else 9)
        if is_bold:
            run.font.color.rgb = _rgb(COLOR_DARK)

    # ==========================================
    # 6. AVANSS UN SUMMA VĀRDIEM
    # ==========================================
    doc.add_paragraph()

    if is_advance:
        raw_advance = data.get('raw_advance', 0.0)
        fmt_advance = fmt_curr(raw_advance)
        pct_val     = int(round(data.get('advance_percent', 0)))

        p_adv = doc.add_paragraph()
        p_adv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p_adv.add_run(f"{tr['advance_payable']} ({pct_val}%): {fmt_advance} €")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = _rgb(COLOR_DARK)

    p_words = doc.add_paragraph()
    p_words.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p_words.add_run(f"{tr['words_prefix']}{data.get('amount_words','')}")
    run.italic = True
    run.font.size = Pt(9)

    # ==========================================
    # 7. KOMENTĀRI
    # ==========================================
    comments = data.get('comments', '').strip()
    if comments:
        doc.add_paragraph()
        _add_thin_line(doc, COLOR_LIGHT)

        p_info = doc.add_paragraph()
        r = p_info.add_run(tr['extra_info'])
        r.bold = True
        r.font.color.rgb = _rgb(COLOR_DARK)

        p_comm = doc.add_paragraph()
        p_comm.add_run(comments)

    # ==========================================
    # 8. PARAKSTI
    # ==========================================
    doc.add_paragraph()
    doc.add_paragraph()
    _add_color_bar(doc, COLOR_MID, height_cm=0.15)
    _add_thin_line(doc, COLOR_LIGHT)
    doc.add_paragraph()

    signatory = data.get('signatory', 'SIA Baltic SEO valdes loceklis Adrians Stankevičs')

    doc_lower = doc_type.lower()
    if "pavadzīme" in doc_lower or "invoice" in doc_lower:
        prep_text = tr['prep_invoice']
        recv_text = tr['recv_invoice']
    elif "avansa" in doc_lower or "advance" in doc_lower:
        prep_text = tr['prep_advance']
        recv_text = tr['recv_advance']
    else:
        prep_text = tr['prep_receipt']
        recv_text = tr['recv_receipt']

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.autofit = False
    sig_table.columns[0].width = Cm(10)
    sig_table.columns[1].width = Cm(7)

    # Rinda 1 — sagatavoja
    cell_00 = sig_table.cell(0, 0)
    p = cell_00.paragraphs[0]
    p.add_run(prep_text)
    run_sig = p.add_run(signatory)
    run_sig.italic = True
    run_sig.font.color.rgb = _rgb(COLOR_DARK)

    cell_01 = sig_table.cell(0, 1)
    p = cell_01.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("__________________________")

    # Rinda 2 — saņēma
    cell_10 = sig_table.cell(1, 0)
    cell_10.paragraphs[0].add_run(recv_text)

    cell_11 = sig_table.cell(1, 1)
    p = cell_11.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("__________________________")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
